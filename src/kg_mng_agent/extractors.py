"""File readers for the data layer.

The agent supports Markdown, plain text, PDF, and DOCX inputs. PDF/DOCX use optional
third-party readers when installed and return actionable errors otherwise.
"""

from __future__ import annotations

import hashlib
import importlib.util
from pathlib import Path

from .models import SourceDocument

SUPPORTED_EXTENSIONS = {".md", ".markdown", ".txt", ".pdf", ".docx"}


class UnsupportedInputError(ValueError):
    """Raised when the CLI receives an unsupported document format."""


def read_document(path: Path) -> SourceDocument:
    """Read and normalize a supported source document."""

    resolved = path.expanduser().resolve()
    if not resolved.exists() or not resolved.is_file():
        raise FileNotFoundError(f"Input file does not exist: {path}")

    extension = resolved.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedInputError(f"Unsupported file type '{extension}'. Supported: {supported}")

    raw_bytes = resolved.read_bytes()
    digest = hashlib.sha256(raw_bytes).hexdigest()

    if extension in {".md", ".markdown", ".txt"}:
        text = raw_bytes.decode("utf-8", errors="replace")
        media_type = "text/markdown" if extension in {".md", ".markdown"} else "text/plain"
    elif extension == ".pdf":
        text = _read_pdf(resolved)
        media_type = "application/pdf"
    else:
        text = _read_docx(resolved)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    normalized = "\n".join(line.rstrip() for line in text.splitlines()).strip()
    if not normalized:
        raise ValueError(f"No extractable text found in: {path}")

    return SourceDocument(path=resolved, media_type=media_type, text=normalized, sha256=digest)


def _read_pdf(path: Path) -> str:
    if importlib.util.find_spec("pypdf") is None:  # pragma: no cover - depends on optional env
        raise RuntimeError("PDF support requires installing the 'pdf' extra: pip install kg-mng-agent[pdf]")

    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _read_docx(path: Path) -> str:
    if importlib.util.find_spec("docx") is None:  # pragma: no cover - depends on optional env
        raise RuntimeError("DOCX support requires installing the 'docx' extra: pip install kg-mng-agent[docx]")

    from docx import Document

    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    return "\n".join(paragraphs + table_cells)
